import os
import tempfile
import zipfile
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .config import (
    A4_LANDSCAPE_HEIGHT_CM,
    A4_LANDSCAPE_WIDTH_CM,
    DEFAULT_FONT_NAME,
    DEFAULT_FONT_SIZE_PT,
    DEFAULT_LINE_SPACING,
    DEFAULT_SPACE_AFTER_PT,
    DEFAULT_SPACE_BEFORE_PT,
    MARGIN_BOTTOM_CM,
    MARGIN_LEFT_CM,
    MARGIN_RIGHT_CM,
    MARGIN_TOP_CM,
    TABLE_COLUMN_WIDTHS_CM,
    TABLE_FONT_NAME,
    TABLE_HEADERS,
)


def sanitize_docx_text(value: str) -> str:
    cleaned_chars = []
    for ch in value:
        code = ord(ch)
        if ch in ('\t', '\n', '\r'):
            cleaned_chars.append(ch)
            continue
        if (
            0x20 <= code <= 0xD7FF
            or 0xE000 <= code <= 0xFFFD
            or 0x10000 <= code <= 0x10FFFF
        ):
            cleaned_chars.append(ch)
    return ''.join(cleaned_chars)


def create_docx(
    output_path: str,
    replicas: List[Dict[str, str]],
    metadata: Dict[str, str],
) -> None:
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    set_page_a4_landscape(section)
    set_default_font(doc, DEFAULT_FONT_NAME)

    for label, value in metadata.items():
        line = sanitize_docx_text(f"{label}: {value}" if value else label)
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(
            paragraph,
            line_spacing=DEFAULT_LINE_SPACING,
            space_before=DEFAULT_SPACE_BEFORE_PT,
            space_after=DEFAULT_SPACE_AFTER_PT,
        )
        run = paragraph.add_run(line)
        set_run_font(
            run,
            DEFAULT_FONT_NAME,
            bold=True,
            size_pt=DEFAULT_FONT_SIZE_PT,
        )

    table = doc.add_table(rows=1 + len(replicas), cols=len(TABLE_HEADERS))
    set_table_layout_fixed(table)
    set_column_widths(table, TABLE_COLUMN_WIDTHS_CM)
    set_table_borders(table)
    set_table_rows_layout(table)
    for col_idx, header in enumerate(TABLE_HEADERS):
        set_cell_text_with_alignment(
            table.rows[0].cells[col_idx],
            header,
            bold=True,
            font_name=TABLE_FONT_NAME,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for idx, replica in enumerate(replicas):
        row = table.rows[idx + 1]
        set_cell_text(row.cells[0], str(idx + 1), bold=True, font_name=TABLE_FONT_NAME)
        set_cell_text(row.cells[1], replica["role"], bold=True, font_name=TABLE_FONT_NAME)
        set_cell_text(row.cells[2], replica["text"], bold=False, font_name=TABLE_FONT_NAME)

    save_docx_safely(doc, output_path)


def validate_docx_structure(path: str) -> None:
    required_entries = {
        '[Content_Types].xml',
        '_rels/.rels',
        'word/document.xml',
    }
    try:
        with zipfile.ZipFile(path, 'r') as archive:
            names = set(archive.namelist())
    except zipfile.BadZipFile as exc:
        raise ValueError(f'Invalid DOCX zip container: {path}') from exc

    missing = sorted(required_entries - names)
    if missing:
        missing_joined = ', '.join(missing)
        raise ValueError(f'DOCX is missing required entries: {missing_joined}')


def save_docx_safely(doc: Document, output_path: str) -> None:
    output = Path(output_path).expanduser()
    output.parent.mkdir(parents=True, exist_ok=True)

    fd, tmp_path = tempfile.mkstemp(
        dir=str(output.parent),
        prefix=f'.{output.stem}.',
        suffix='.docx',
    )
    os.close(fd)

    try:
        doc.save(tmp_path)
        validate_docx_structure(tmp_path)
        os.replace(tmp_path, output)
    except Exception:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        raise


def set_default_font(doc: Document, font_name: str) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(DEFAULT_FONT_SIZE_PT)
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)


def set_run_font(
    run,
    font_name: str,
    bold: Optional[bool] = None,
    size_pt: Optional[int] = None,
) -> None:
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    if bold is not None:
        run.bold = bold


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    font_name: str = DEFAULT_FONT_NAME,
) -> None:
    cell.text = sanitize_docx_text(text)
    if not cell.paragraphs:
        return
    for paragraph in cell.paragraphs:
        set_paragraph_spacing(
            paragraph,
            line_spacing=DEFAULT_LINE_SPACING,
            space_before=DEFAULT_SPACE_BEFORE_PT,
            space_after=DEFAULT_SPACE_AFTER_PT,
        )
        for run in paragraph.runs:
            set_run_font(
                run,
                font_name,
                bold=bold,
                size_pt=DEFAULT_FONT_SIZE_PT,
            )


def set_cell_text_with_alignment(
    cell,
    text: str,
    bold: bool,
    font_name: str,
    alignment: WD_ALIGN_PARAGRAPH,
) -> None:
    set_cell_text(cell, text, bold=bold, font_name=font_name)
    for paragraph in cell.paragraphs:
        paragraph.alignment = alignment


def set_column_widths(table, widths_cm: List[float]) -> None:
    table.autofit = False
    if hasattr(table, "allow_autofit"):
        table.allow_autofit = False
    total_width = sum(widths_cm)
    set_table_width(table, total_width)
    set_table_grid(table, widths_cm)
    for idx, width in enumerate(widths_cm):
        table.columns[idx].width = Cm(width)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            set_cell_width(row.cells[idx], width)


def set_cell_width(cell, width_cm: float) -> None:
    width = Cm(width_cm)
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width.twips))


def set_table_layout_fixed(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(Cm(width_cm).twips))


def set_table_grid(table, widths_cm: List[float]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(tbl.index(tbl_pr) + 1, tbl_grid)
    else:
        tbl_grid.clear()
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(Cm(width).twips))
        tbl_grid.append(grid_col)


def set_paragraph_spacing(
    paragraph,
    line_spacing: float,
    space_before: int,
    space_after: int,
) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = line_spacing
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)


def set_table_borders(table, size: int = 4) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(size))
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "auto")
        borders.append(elem)
    tbl_pr.append(borders)


def set_page_a4_landscape(section) -> None:
    section.page_width = Cm(A4_LANDSCAPE_WIDTH_CM)
    section.page_height = Cm(A4_LANDSCAPE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_TOP_CM)
    section.bottom_margin = Cm(MARGIN_BOTTOM_CM)
    section.left_margin = Cm(MARGIN_LEFT_CM)
    section.right_margin = Cm(MARGIN_RIGHT_CM)


def set_table_rows_layout(table, min_height_pt: Optional[int] = None) -> None:
    for row in table.rows:
        row.allow_break_across_pages = True
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        if min_height_pt is not None:
            row.height = Pt(min_height_pt)
